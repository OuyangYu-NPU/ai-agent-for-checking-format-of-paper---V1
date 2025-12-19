"""
PDF转换工具
提取PDF中的文字、格式、图片，保存为Word、Markdown或JSON文档
"""

import fitz  # PyMuPDF
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import io
import os
import sys
import json
import base64
from pathlib import Path
from typing import List, Dict, Any, Optional


class PDFConverter:
    """PDF转换器基类"""
    
    def __init__(self, pdf_path: str):
        """
        初始化转换器
        
        Args:
            pdf_path: PDF文件路径
        """
        self.pdf_path = Path(pdf_path)
        if not self.pdf_path.exists():
            raise FileNotFoundError(f"PDF文件不存在: {pdf_path}")
        
        self.pdf_doc = fitz.open(str(self.pdf_path))
    
    def _get_font_size(self, span: dict) -> float:
        """获取字体大小"""
        return span.get('size', 11)
    
    def _is_bold(self, span: dict) -> bool:
        """判断是否粗体"""
        flags = span.get('flags', 0)
        return bool(flags & 2 ** 4)
    
    def _is_italic(self, span: dict) -> bool:
        """判断是否斜体"""
        flags = span.get('flags', 0)
        return bool(flags & 2 ** 1)
    
    def _get_text_color(self, span: dict) -> tuple:
        """获取文字颜色"""
        color = span.get('color', 0)
        if isinstance(color, int):
            r = (color >> 16) & 0xFF
            g = (color >> 8) & 0xFF
            b = color & 0xFF
            return (r, g, b)
        return (0, 0, 0)
    
    def _extract_images_from_page(self, page) -> list:
        """从页面提取图片"""
        images = []
        image_list = page.get_images(full=True)
        
        for img_info in image_list:
            xref = img_info[0]
            try:
                base_image = self.pdf_doc.extract_image(xref)
                image_bytes = base_image["image"]
                img_rects = page.get_image_rects(xref)
                if img_rects:
                    rect = img_rects[0]
                    images.append((image_bytes, rect, base_image.get("ext", "png")))
            except Exception as e:
                print(f"提取图片时出错: {e}")
                continue
        
        return images
    
    def close(self):
        """关闭PDF文档"""
        if self.pdf_doc:
            self.pdf_doc.close()
    
    def __enter__(self):
        return self
    
    def __exit__(self, exc_type, exc_val, exc_tb):
        self.close()


class PDFToWordConverter(PDFConverter):
    """PDF转Word转换器"""
    
    def __init__(self, pdf_path: str):
        super().__init__(pdf_path)
        self.doc = Document()
        self._set_default_font()
    
    def _set_default_font(self):
        """设置文档默认字体"""
        style = self.doc.styles['Normal']
        style.font.name = '微软雅黑'
        style.font.size = Pt(11)
        style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    
    def _add_image_to_doc(self, image_bytes: bytes, width_inches: float = None):
        """添加图片到Word文档"""
        try:
            image_stream = io.BytesIO(image_bytes)
            if width_inches is None or width_inches > 6:
                width_inches = 6
            
            paragraph = self.doc.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run()
            run.add_picture(image_stream, width=Inches(width_inches))
        except Exception as e:
            print(f"添加图片到Word时出错: {e}")
    
    def _process_text_block(self, block: dict):
        """处理文本块"""
        if "lines" not in block:
            return
        
        paragraph = self.doc.add_paragraph()
        
        for line in block["lines"]:
            for span in line.get("spans", []):
                text = span.get("text", "")
                if not text.strip():
                    continue
                
                run = paragraph.add_run(text)
                font_size = self._get_font_size(span)
                run.font.size = Pt(font_size)
                run.font.bold = self._is_bold(span)
                run.font.italic = self._is_italic(span)
                
                r, g, b = self._get_text_color(span)
                if (r, g, b) != (0, 0, 0):
                    run.font.color.rgb = RGBColor(r, g, b)
                
                font_name = span.get("font", "")
                if font_name:
                    run.font.name = font_name
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    
    def convert(self, output_path: str = None, progress_callback=None) -> str:
        """执行转换"""
        if output_path is None:
            output_path = str(self.pdf_path.with_suffix('.docx'))
        
        total_pages = len(self.pdf_doc)
        print(f"开始转换为Word，共 {total_pages} 页...")
        
        for page_num in range(total_pages):
            page = self.pdf_doc[page_num]
            blocks = page.get_text("dict", flags=fitz.TEXT_PRESERVE_WHITESPACE)["blocks"]
            images = self._extract_images_from_page(page)
            
            elements = []
            for block in blocks:
                if block["type"] == 0:
                    y_pos = block["bbox"][1]
                    elements.append(("text", y_pos, block))
            
            for img_bytes, rect, ext in images:
                y_pos = rect.y0
                width_inches = rect.width / 72
                elements.append(("image", y_pos, (img_bytes, width_inches)))
            
            elements.sort(key=lambda x: x[1])
            
            for elem_type, _, data in elements:
                if elem_type == "text":
                    self._process_text_block(data)
                elif elem_type == "image":
                    img_bytes, width = data
                    self._add_image_to_doc(img_bytes, width)
            
            if page_num < total_pages - 1:
                self.doc.add_page_break()
            
            if progress_callback:
                progress_callback(page_num + 1, total_pages)
            
            print(f"已处理第 {page_num + 1}/{total_pages} 页")
        
        self.doc.save(output_path)
        print(f"转换完成！已保存到: {output_path}")
        
        return output_path


class PDFToMarkdownConverter(PDFConverter):
    """PDF转Markdown转换器"""
    
    def __init__(self, pdf_path: str, save_images: bool = True):
        """
        初始化Markdown转换器
        
        Args:
            pdf_path: PDF文件路径
            save_images: 是否保存图片到文件（False则使用base64内嵌）
        """
        super().__init__(pdf_path)
        self.save_images = save_images
        self.image_dir = None
        self.image_count = 0
        self.md_content = []
    
    def _detect_heading_level(self, font_size: float) -> int:
        """根据字体大小检测标题级别"""
        if font_size >= 24:
            return 1
        elif font_size >= 20:
            return 2
        elif font_size >= 16:
            return 3
        elif font_size >= 14:
            return 4
        elif font_size >= 12:
            return 5
        return 0  # 普通文本
    
    def _format_text(self, text: str, is_bold: bool, is_italic: bool) -> str:
        """格式化文本（加粗、斜体）"""
        if not text.strip():
            return text
        
        if is_bold and is_italic:
            return f"***{text}***"
        elif is_bold:
            return f"**{text}**"
        elif is_italic:
            return f"*{text}*"
        return text
    
    def _save_image(self, image_bytes: bytes, ext: str) -> str:
        """保存图片并返回引用路径"""
        self.image_count += 1
        
        if self.save_images and self.image_dir:
            # 保存到文件
            image_filename = f"image_{self.image_count}.{ext}"
            image_path = self.image_dir / image_filename
            with open(image_path, 'wb') as f:
                f.write(image_bytes)
            return f"./images/{image_filename}"
        else:
            # 使用base64内嵌
            b64_data = base64.b64encode(image_bytes).decode('utf-8')
            mime_type = f"image/{ext}" if ext != "jpg" else "image/jpeg"
            return f"data:{mime_type};base64,{b64_data}"
    
    def _process_text_block_md(self, block: dict) -> str:
        """处理文本块并返回Markdown文本"""
        if "lines" not in block:
            return ""
        
        lines_text = []
        current_line_parts = []
        max_font_size = 0
        has_bold = False
        
        for line in block["lines"]:
            for span in line.get("spans", []):
                text = span.get("text", "")
                if not text.strip():
                    continue
                
                font_size = self._get_font_size(span)
                is_bold = self._is_bold(span)
                is_italic = self._is_italic(span)
                
                max_font_size = max(max_font_size, font_size)
                if is_bold:
                    has_bold = True
                
                formatted_text = self._format_text(text, is_bold, is_italic)
                current_line_parts.append(formatted_text)
        
        if not current_line_parts:
            return ""
        
        line_text = "".join(current_line_parts)
        
        # 检测是否为标题
        heading_level = self._detect_heading_level(max_font_size)
        if heading_level > 0 and has_bold:
            # 移除格式标记用于标题
            clean_text = line_text.replace("**", "").replace("*", "")
            return f"{'#' * heading_level} {clean_text}"
        
        return line_text
    
    def convert(self, output_path: str = None, progress_callback=None) -> str:
        """执行转换"""
        if output_path is None:
            output_path = str(self.pdf_path.with_suffix('.md'))
        
        output_path = Path(output_path)
        
        # 创建图片目录
        if self.save_images:
            self.image_dir = output_path.parent / "images"
            self.image_dir.mkdir(exist_ok=True)
        
        total_pages = len(self.pdf_doc)
        print(f"开始转换为Markdown，共 {total_pages} 页...")
        
        self.md_content = []
        
        for page_num in range(total_pages):
            page = self.pdf_doc[page_num]
            blocks = page.get_text("dict", flags=fitz.TEXT_PRESERVE_WHITESPACE)["blocks"]
            images = self._extract_images_from_page(page)
            
            elements = []
            for block in blocks:
                if block["type"] == 0:
                    y_pos = block["bbox"][1]
                    elements.append(("text", y_pos, block))
            
            for img_bytes, rect, ext in images:
                y_pos = rect.y0
                elements.append(("image", y_pos, (img_bytes, ext)))
            
            elements.sort(key=lambda x: x[1])
            
            page_content = []
            for elem_type, _, data in elements:
                if elem_type == "text":
                    text = self._process_text_block_md(data)
                    if text.strip():
                        page_content.append(text)
                elif elem_type == "image":
                    img_bytes, ext = data
                    img_ref = self._save_image(img_bytes, ext)
                    page_content.append(f"\n![图片]({img_ref})\n")
            
            if page_content:
                self.md_content.append("\n\n".join(page_content))
            
            # 添加分页标记
            if page_num < total_pages - 1:
                self.md_content.append("\n\n---\n\n")
            
            if progress_callback:
                progress_callback(page_num + 1, total_pages)
            
            print(f"已处理第 {page_num + 1}/{total_pages} 页")
        
        # 保存Markdown文件
        final_content = "".join(self.md_content)
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(final_content)
        
        print(f"转换完成！已保存到: {output_path}")
        
        return str(output_path)


class PDFToJSONConverter(PDFConverter):
    """PDF转JSON转换器 - 保留页码、页眉页脚信息，文本采用Markdown格式"""
    
    def __init__(self, pdf_path: str, header_ratio: float = 0.06, footer_ratio: float = 0.05):
        """
        初始化JSON转换器
        
        Args:
            pdf_path: PDF文件路径
            header_ratio: 页眉区域占页面高度的比例（默认6%，约50点/1.8cm）
            footer_ratio: 页脚区域占页面高度的比例（默认5%，约42点/1.5cm）
        """
        super().__init__(pdf_path)
        self.header_ratio = header_ratio
        self.footer_ratio = footer_ratio
        self.image_count = 0
    
    def _detect_heading_level(self, font_size: float) -> int:
        """根据字体大小检测标题级别"""
        if font_size >= 24:
            return 1
        elif font_size >= 20:
            return 2
        elif font_size >= 16:
            return 3
        elif font_size >= 14:
            return 4
        elif font_size >= 12:
            return 5
        return 0  # 普通文本
    
    def _format_text_md(self, text: str, is_bold: bool, is_italic: bool) -> str:
        """格式化文本为Markdown格式"""
        if not text.strip():
            return text
        
        if is_bold and is_italic:
            return f"***{text}***"
        elif is_bold:
            return f"**{text}**"
        elif is_italic:
            return f"*{text}*"
        return text
    
    def _classify_block_position(self, block: dict, page_height: float) -> str:
        """
        根据位置分类文本块
        
        Args:
            block: 文本块
            page_height: 页面高度
            
        Returns:
            位置类型: "header", "footer", "body"
        """
        y_top = block["bbox"][1]
        y_bottom = block["bbox"][3]
        block_height = y_bottom - y_top
        
        header_threshold = page_height * self.header_ratio
        footer_threshold = page_height * (1 - self.footer_ratio)
        
        # 获取文本块的文本内容长度（用于辅助判断）
        text_length = 0
        if "lines" in block:
            for line in block["lines"]:
                for span in line.get("spans", []):
                    text_length += len(span.get("text", ""))
        
        # 页眉判断：顶部位置在页眉区域内
        if y_top < header_threshold:
            return "header"
        
        # 页脚判断：更严格的条件
        # 1. 底部位置在页脚区域内
        # 2. 文本块高度较小（单行或少量行）
        # 3. 或者文本内容很短（通常是页码）
        if y_bottom > footer_threshold:
            # 额外检查：正常页脚通常高度较小且内容较短
            is_small_block = block_height < page_height * 0.03  # 高度小于3%
            is_short_text = text_length < 30  # 内容少于30字符
            
            if is_small_block or is_short_text:
                return "footer"
        
        return "body"
    
    def _extract_lines_from_block(self, block: dict) -> List[Dict[str, Any]]:
        """
        从文本块中提取每一行作为独立元素
        基于 span 的 bbox y坐标来精确区分行（不依赖 PyMuPDF 的 line 结构）
        
        Args:
            block: 文本块
            
        Returns:
            行信息列表，每行一个字典
        """
        if "lines" not in block:
            return []
        
        # 收集所有 spans，按 y 坐标聚类成行
        all_spans = []
        for line in block["lines"]:
            for span in line.get("spans", []):
                if "bbox" in span and span.get("text"):
                    all_spans.append(span)
        
        if not all_spans:
            return []
        
        # 按 y0 坐标排序
        all_spans.sort(key=lambda s: (s["bbox"][1], s["bbox"][0]))
        
        # 聚类：y坐标差距超过阈值则视为不同行
        # 阈值设为字体高度的一半
        clustered_lines = []
        current_line_spans = [all_spans[0]]
        current_y_center = (all_spans[0]["bbox"][1] + all_spans[0]["bbox"][3]) / 2
        
        for span in all_spans[1:]:
            span_y_center = (span["bbox"][1] + span["bbox"][3]) / 2
            span_height = span["bbox"][3] - span["bbox"][1]
            
            # 如果 y 中心差距超过行高的 60%，视为新行
            threshold = max(span_height * 0.6, 5)  # 至少5个点
            
            if abs(span_y_center - current_y_center) > threshold:
                # 新的一行
                clustered_lines.append(current_line_spans)
                current_line_spans = [span]
                current_y_center = span_y_center
            else:
                current_line_spans.append(span)
                # 更新当前行的 y 中心（取平均值）
                y_centers = [(s["bbox"][1] + s["bbox"][3]) / 2 for s in current_line_spans]
                current_y_center = sum(y_centers) / len(y_centers)
        
        # 添加最后一行
        if current_line_spans:
            clustered_lines.append(current_line_spans)
        
        # 构建行信息
        lines_result = []
        for line_spans in clustered_lines:
            # 计算行的 bbox
            line_bbox = [
                min(s["bbox"][0] for s in line_spans),
                min(s["bbox"][1] for s in line_spans),
                max(s["bbox"][2] for s in line_spans),
                max(s["bbox"][3] for s in line_spans)
            ]
            
            raw_text_parts = []
            md_text_parts = []
            spans_info = []
            max_font_size = 0
            has_bold = False
            
            # 按 x 坐标排序 spans
            line_spans.sort(key=lambda s: s["bbox"][0])
            
            for span in line_spans:
                text = span.get("text", "")
                if not text:
                    continue
                
                font_size = self._get_font_size(span)
                is_bold = self._is_bold(span)
                is_italic = self._is_italic(span)
                
                max_font_size = max(max_font_size, font_size)
                if is_bold:
                    has_bold = True
                
                span_info = {
                    "text": text,
                    "font": span.get("font", ""),
                    "size": round(font_size, 2),
                    "bold": is_bold,
                    "italic": is_italic,
                    "color": self._get_text_color(span)
                }
                spans_info.append(span_info)
                raw_text_parts.append(text)
                
                md_text = self._format_text_md(text, is_bold, is_italic)
                md_text_parts.append(md_text)
            
            if not raw_text_parts:
                continue
            
            # 生成Markdown格式内容
            md_content = "".join(md_text_parts)
            raw_text = "".join(raw_text_parts)
            
            # 检测是否为标题
            heading_level = self._detect_heading_level(max_font_size)
            if heading_level > 0 and has_bold:
                clean_text = md_content.replace("**", "").replace("*", "")
                md_content = f"{'#' * heading_level} {clean_text}"
            
            line_info = {
                "type": "text",
                "content": md_content,
                "raw_text": raw_text,
                "is_heading": heading_level > 0 and has_bold,
                "heading_level": heading_level if (heading_level > 0 and has_bold) else None,
                "bbox": {
                    "x0": round(line_bbox[0], 2),
                    "y0": round(line_bbox[1], 2),
                    "x1": round(line_bbox[2], 2),
                    "y1": round(line_bbox[3], 2)
                },
                "spans": spans_info
            }
            lines_result.append(line_info)
        
        return lines_result
    
    def _classify_line_position(self, line_info: dict, page_height: float) -> str:
        """
        根据行的bbox位置分类
        
        Args:
            line_info: 行信息（包含bbox）
            page_height: 页面高度
            
        Returns:
            位置类型: "header", "footer", "body"
        """
        y_top = line_info["bbox"]["y0"]
        y_bottom = line_info["bbox"]["y1"]
        line_height = y_bottom - y_top
        raw_text = line_info.get("raw_text", "").strip()
        text_length = len(raw_text)
        
        header_threshold = page_height * self.header_ratio
        # 页脚区域：页面底部 8% 的区域（约67点/2.4cm for A4）
        footer_threshold = page_height * 0.92
        
        # 页眉判断
        if y_top < header_threshold:
            return "header"
        
        # 页脚判断：使用 y_top（行顶部）超过阈值
        if y_top > footer_threshold:
            return "footer"
        
        # 额外检查：如果内容像页码（短文本，在页面下方），也视为页脚
        # 页码特征：内容很短，位置在页面下方 85% 以下
        if y_top > page_height * 0.85:
            is_short_text = text_length < 15
            is_page_number_like = self._looks_like_page_number(raw_text)
            if is_short_text or is_page_number_like:
                return "footer"
        
        return "body"
    
    def _looks_like_page_number(self, text: str) -> bool:
        """
        判断文本是否像页码
        
        支持格式：
        - 阿拉伯数字: 1, 2, 3
        - 罗马数字: I, II, III, IV, V, i, ii, iii
        - 带装饰: - 1 -, — 1 —, [ 1 ], ( 1 )
        - 中文格式: 第1页, 第一页
        """
        import re
        
        text = text.strip()
        if not text:
            return False
        
        # 罗马数字模式
        roman_pattern = r'^[\s\-—\[\]()]*[IVXLCDMivxlcdm]+[\s\-—\[\]()]*$'
        
        # 阿拉伯数字模式
        arabic_pattern = r'^[\s\-—\[\]()]*\d+[\s\-—\[\]()]*$'
        
        # 分数格式 1/10
        fraction_pattern = r'^\d+\s*/\s*\d+$'
        
        # 中文格式
        chinese_pattern = r'^第\s*[\d一二三四五六七八九十百]+\s*页$'
        
        # Page 格式
        page_pattern = r'^page\s*\d+$'
        
        for pattern in [roman_pattern, arabic_pattern, fraction_pattern, chinese_pattern, page_pattern]:
            if re.match(pattern, text, re.IGNORECASE):
                return True
        
        return False
    
    def _extract_image_info(self, image_bytes: bytes, rect, ext: str, 
                            save_dir: Optional[Path] = None) -> Dict[str, Any]:
        """
        提取图片信息
        
        Args:
            image_bytes: 图片字节数据
            rect: 图片位置矩形
            ext: 图片扩展名
            save_dir: 图片保存目录
            
        Returns:
            图片信息字典
        """
        self.image_count += 1
        
        # 保存图片或转base64
        if save_dir:
            image_filename = f"image_{self.image_count}.{ext}"
            image_path = save_dir / image_filename
            with open(image_path, 'wb') as f:
                f.write(image_bytes)
            image_ref = f"./images/{image_filename}"
        else:
            b64_data = base64.b64encode(image_bytes).decode('utf-8')
            mime_type = f"image/{ext}" if ext != "jpg" else "image/jpeg"
            image_ref = f"data:{mime_type};base64,{b64_data}"
        
        return {
            "type": "image",
            "src": image_ref,
            "markdown": f"![图片{self.image_count}]({image_ref})",
            "format": ext,
            "bbox": {
                "x0": round(rect.x0, 2),
                "y0": round(rect.y0, 2),
                "x1": round(rect.x1, 2),
                "y1": round(rect.y1, 2)
            },
            "width": round(rect.width, 2),
            "height": round(rect.height, 2)
        }
    
    def _extract_page_number_from_footer(self, footer_text: str) -> Optional[str]:
        """
        从页脚文本中提取页码
        
        支持的格式：
        - "1", "2", "3" (纯数字)
        - "I", "II", "III", "IV" (罗马数字)
        - "- 1 -", "- 2 -" (带横线)
        - "第1页", "第2页" (中文格式)
        - "Page 1", "page 2" (英文格式)
        - "1/10", "2/10" (分数格式)
        
        Args:
            footer_text: 页脚文本
            
        Returns:
            提取的页码字符串，未找到返回None
        """
        import re
        
        if not footer_text:
            return None
        
        footer_text = footer_text.strip()
        
        # 尝试各种页码格式
        patterns = [
            (r'第\s*(\d+)\s*页', 1),                    # 第1页
            (r'Page\s*(\d+)', 1),                       # Page 1
            (r'-\s*(\d+)\s*-', 1),                      # - 1 -
            (r'—\s*(\d+)\s*—', 1),                      # — 1 —
            (r'(\d+)\s*/\s*\d+', 1),                    # 1/10
            (r'^\s*(\d+)\s*$', 1),                      # 纯数字
            (r'^\s*([IVXLCDMivxlcdm]+)\s*$', 1),       # 纯罗马数字
            (r'-\s*([IVXLCDMivxlcdm]+)\s*-', 1),       # - I -
            (r'—\s*([IVXLCDMivxlcdm]+)\s*—', 1),       # — I —
        ]
        
        for pattern, group in patterns:
            match = re.search(pattern, footer_text, re.IGNORECASE)
            if match:
                return match.group(group)
        
        return None
    
    def convert(self, output_path: str = None, progress_callback=None, 
                save_images: bool = True) -> str:
        """
        执行转换
        
        Args:
            output_path: 输出JSON文件路径
            progress_callback: 进度回调函数
            save_images: 是否保存图片到文件
            
        Returns:
            输出文件路径
        """
        if output_path is None:
            output_path = str(self.pdf_path.with_suffix('.json'))
        
        output_path = Path(output_path)
        
        # 创建图片目录
        image_dir = None
        if save_images:
            image_dir = output_path.parent / "images"
            image_dir.mkdir(exist_ok=True)
        
        total_pages = len(self.pdf_doc)
        print(f"开始转换为JSON，共 {total_pages} 页...")
        
        # 文档级别信息
        result = {
            "source": str(self.pdf_path.name),
            "total_pages": total_pages,
            "metadata": {
                "title": self.pdf_doc.metadata.get("title", ""),
                "author": self.pdf_doc.metadata.get("author", ""),
                "subject": self.pdf_doc.metadata.get("subject", ""),
                "creator": self.pdf_doc.metadata.get("creator", ""),
                "producer": self.pdf_doc.metadata.get("producer", ""),
                "creation_date": self.pdf_doc.metadata.get("creationDate", ""),
                "modification_date": self.pdf_doc.metadata.get("modDate", "")
            },
            "pages": []
        }
        
        for page_num in range(total_pages):
            page = self.pdf_doc[page_num]
            page_rect = page.rect
            page_height = page_rect.height
            page_width = page_rect.width
            
            # 获取文本块
            blocks = page.get_text("dict", flags=fitz.TEXT_PRESERVE_WHITESPACE)["blocks"]
            
            # 提取图片
            images = self._extract_images_from_page(page)
            
            # 分类内容
            header_content = []
            body_content = []
            footer_content = []
            
            # 处理文本块 - 按行级别提取和分类
            for block in blocks:
                if block["type"] == 0:  # 文本块
                    # 提取块中的每一行
                    lines = self._extract_lines_from_block(block)
                    
                    for line_info in lines:
                        # 根据每行的bbox独立判断位置
                        position = self._classify_line_position(line_info, page_height)
                        line_info["position"] = position
                        
                        if position == "header":
                            header_content.append(line_info)
                        elif position == "footer":
                            footer_content.append(line_info)
                        else:
                            body_content.append(line_info)
            
            # 处理图片
            for img_bytes, rect, ext in images:
                img_info = self._extract_image_info(img_bytes, rect, ext, image_dir)
                
                # 根据位置分类图片（图片一般不会在页脚，判断更严格）
                if rect.y0 < page_height * self.header_ratio:
                    img_info["position"] = "header"
                    header_content.append(img_info)
                else:
                    # 图片很少出现在页脚，除非是非常小的logo
                    # 只有当图片完全在页脚区域且很小时才算页脚
                    is_in_footer_zone = rect.y0 > page_height * (1 - self.footer_ratio)
                    is_small_image = rect.height < page_height * 0.03
                    
                    if is_in_footer_zone and is_small_image:
                        img_info["position"] = "footer"
                        footer_content.append(img_info)
                    else:
                        img_info["position"] = "body"
                        body_content.append(img_info)
            
            # 按y坐标排序
            header_content.sort(key=lambda x: x["bbox"]["y0"])
            body_content.sort(key=lambda x: x["bbox"]["y0"])
            footer_content.sort(key=lambda x: x["bbox"]["y0"])
            
            # 辅助函数：生成包含图片的Markdown内容
            def build_markdown_with_images(elements: List[Dict]) -> str:
                md_parts = []
                for item in elements:
                    if item["type"] == "text":
                        md_parts.append(item["content"])
                    elif item["type"] == "image":
                        md_parts.append(item["markdown"])
                return "\n\n".join(md_parts)
            
            # 提取页脚原始文本和页码
            footer_raw_text = " ".join([
                item.get("raw_text", "") for item in footer_content 
                if item["type"] == "text"
            ])
            footer_page_number = self._extract_page_number_from_footer(footer_raw_text)
            
            # 统计该页所有字体
            all_elements = header_content + body_content + footer_content
            page_fonts = set()
            for elem in all_elements:
                if elem["type"] == "text":
                    for span in elem.get("spans", []):
                        font = span.get("font", "")
                        if font:
                            page_fonts.add(font)
            
            # 提取图片注释（图片下方紧邻的文字行）
            image_captions = []
            all_sorted_elements = sorted(all_elements, key=lambda x: x["bbox"]["y0"])
            
            for i, elem in enumerate(all_sorted_elements):
                if elem["type"] == "image":
                    img_bottom = elem["bbox"]["y1"]
                    img_src = elem.get("src", "")
                    
                    # 找图片下方最近的文本行
                    caption_text = None
                    caption_distance = float('inf')
                    
                    for j in range(i + 1, len(all_sorted_elements)):
                        next_elem = all_sorted_elements[j]
                        if next_elem["type"] == "text":
                            text_top = next_elem["bbox"]["y0"]
                            distance = text_top - img_bottom
                            
                            # 图片下方 50 点以内的文本视为图注释
                            if 0 < distance < 50 and distance < caption_distance:
                                caption_text = next_elem.get("raw_text", "").strip()
                                caption_distance = distance
                                break  # 取最近的一行
                    
                    if caption_text:
                        image_captions.append({
                            "image_src": img_src,
                            "caption": caption_text,
                            "image_bbox": elem["bbox"]
                        })
            
            # 构建页面数据（使用Markdown格式组织文本和图片）
            page_data = {
                "page_number": page_num + 1,
                "footer_page_number": footer_page_number,
                "width": round(page_width, 2),
                "height": round(page_height, 2),
                "fonts": sorted(list(page_fonts)),
                "image_captions": image_captions if image_captions else None,
                "header": {
                    "markdown": build_markdown_with_images(header_content),
                    "raw_text": " ".join([
                        item.get("raw_text", "") for item in header_content 
                        if item["type"] == "text"
                    ]),
                    "elements": header_content
                },
                "body": {
                    "markdown": build_markdown_with_images(body_content),
                    "raw_text": " ".join([
                        item.get("raw_text", "") for item in body_content 
                        if item["type"] == "text"
                    ]),
                    "elements": body_content
                },
                "footer": {
                    "markdown": build_markdown_with_images(footer_content),
                    "raw_text": footer_raw_text,
                    "detected_page_number": footer_page_number,
                    "elements": footer_content
                }
            }
            
            result["pages"].append(page_data)
            
            if progress_callback:
                progress_callback(page_num + 1, total_pages)
            
            print(f"已处理第 {page_num + 1}/{total_pages} 页")
        
        # 保存JSON文件
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(result, f, ensure_ascii=False, indent=2)
        
        print(f"转换完成！已保存到: {output_path}")
        
        return str(output_path)


def convert_pdf(pdf_path: str, output_path: str = None, output_format: str = "docx",
                progress_callback=None) -> str:
    """
    将PDF转换为指定格式
    
    Args:
        pdf_path: PDF文件路径
        output_path: 输出文件路径，None则自动生成
        output_format: 输出格式 "docx", "md" 或 "json"
        progress_callback: 进度回调函数
        
    Returns:
        输出文件路径
    """
    format_lower = output_format.lower()
    
    if format_lower in ("md", "markdown"):
        with PDFToMarkdownConverter(pdf_path) as converter:
            return converter.convert(output_path, progress_callback)
    elif format_lower == "json":
        with PDFToJSONConverter(pdf_path) as converter:
            return converter.convert(output_path, progress_callback)
    else:
        with PDFToWordConverter(pdf_path) as converter:
            return converter.convert(output_path, progress_callback)


def main():
    """主函数 - 命令行入口"""
    if len(sys.argv) < 2:
        print("用法: python pdf_converter.py <PDF文件路径> [输出路径] [--format docx|md|json]")
        print("示例:")
        print("  python pdf_converter.py input.pdf")
        print("  python pdf_converter.py input.pdf output.docx")
        print("  python pdf_converter.py input.pdf output.md --format md")
        print("  python pdf_converter.py input.pdf output.json --format json")
        print("\n支持的格式:")
        print("  docx - Word文档（默认）")
        print("  md   - Markdown文档")
        print("  json - JSON格式（保留页码、页眉页脚信息）")
        sys.exit(1)
    
    pdf_path = sys.argv[1]
    output_path = None
    output_format = "docx"
    
    # 解析参数
    args = sys.argv[2:]
    i = 0
    while i < len(args):
        if args[i] == "--format" and i + 1 < len(args):
            output_format = args[i + 1]
            i += 2
        else:
            output_path = args[i]
            i += 1
    
    try:
        result_path = convert_pdf(pdf_path, output_path, output_format)
        print(f"\n✓ 转换成功！")
        print(f"  输出文件: {result_path}")
    except FileNotFoundError as e:
        print(f"✗ 错误: {e}")
        sys.exit(1)
    except Exception as e:
        print(f"✗ 转换失败: {e}")
        sys.exit(1)


if __name__ == "__main__":
    main()

